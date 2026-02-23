# -*- coding: utf-8 -*-
from flask import Flask, render_template, abort, request, jsonify, send_file, Response 
import json
import os
import io
import time
import sys 
import logging # 터미널 로그 제어용
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv
from openai import OpenAI

# [로깅 봉인] 터미널에 한글 로그가 찍혀서 서버가 죽는 현상을 완벽히 차단합니다.
logging.getLogger('werkzeug').setLevel(logging.ERROR)

# [핵심 추가] debug=True 일 때 OpenAI 내부 모듈(httpx)이 한글 데이터를 
# 터미널에 몰래 출력하다가 서버를 터뜨리는(Connection error) 현상을 원천 차단합니다!
logging.getLogger('httpx').setLevel(logging.WARNING)
logging.getLogger('openai').setLevel(logging.WARNING)

# 환경 변수 로드
load_dotenv()
DEFAULT_API_KEY = os.getenv("OPENAI_API_KEY")

app = Flask(__name__)
app.jinja_env.add_extension('jinja2.ext.do')

# JSON 인코딩 설정 (False로 설정하여 유니코드 유지)
app.config['JSON_AS_ASCII'] = False 
app.config['JSONIFY_PRETTYPRINT_REGULAR'] = True

def load_data():
    """전술 데이터베이스(content.json) 로드 엔진
       -> 확장된 30여 개의 데이터도 알아서 100% 동적 로드합니다.
    """
    with open('content.json', 'r', encoding='utf-8') as f:
        return json.load(f)

@app.route('/')
def index():
    data = load_data()
    return render_template('index.html', patterns=data.get('prompt_patterns', []))

@app.route('/detail/<int:pattern_id>')
def detail(pattern_id):
    data = load_data()
    patterns = data.get('prompt_patterns', [])
    
    # 확장된 JSON 데이터 안에서 현재 ID를 동적으로 찾아냅니다.
    pattern_index = next((i for i, p in enumerate(patterns) if p['id'] == pattern_id), None)
    if pattern_index is None: abort(404)
    pattern = patterns[pattern_index]
    
    # 26번(Lab) 모듈일 경우의 분기 처리 완벽 유지
    if pattern_id == 26:
        return render_template('lab.html', patterns=patterns)
    
    # 다음 챕터로 넘어가는 로직 (데이터가 추가되어도 자동으로 다음 내용을 물고 옴)
    next_pattern = patterns[pattern_index + 1] if pattern_index + 1 < len(patterns) else None
    
    return render_template('detail.html', 
                            pattern=pattern, 
                            patterns=patterns, 
                            next_pattern=next_pattern)

# ---------------------------------------------------------
# 4. [엔진 보강] 지고의 프롬프트 생성 엔진 (바이트 전송 및 키 검증 강화)
# ---------------------------------------------------------
@app.route('/generate_tactic', methods=['POST'])
def generate_tactic():
    data = request.json
    user_purpose = data.get('purpose')
    user_api_key = data.get('api_key')

    if not user_api_key or len(user_api_key.strip()) < 10:
        err_json = json.dumps({"error": "사이드바에서 API Key를 먼저 동기화(SAVE)해주세요."}, ensure_ascii=False).encode('utf-8')
        return Response(err_json, status=401, mimetype='application/json; charset=utf-8')

    if not user_purpose: 
        err_json = json.dumps({"error": "목적 데이터 누락"}, ensure_ascii=False).encode('utf-8')
        return Response(err_json, status=400, mimetype='application/json; charset=utf-8')

    try:
        active_client = OpenAI(api_key=user_api_key)

        # [원칙 준수] 대표님의 시스템 프롬프트 멘트 100% 보존
        system_instruction = """
        당신은 세계 최고의 'Supreme Prompt Architect'입니다. 
        사용자의 요청을 분석하여 실무에서 즉시 '압도적인 성과'를 내는 초격차 프롬프트를 설계합니다.

        [프롬프트 설계 5대 구성 요소]
        1. [역할(Persona)]: 해당 분야의 1% 전문가로서의 정체성을 부여할 것.
        2. [배경과 의도, 맥락(Context)]: 사용자의 숨은 의도와 비즈니스 환경을 깊이 있게 통찰할 것.
        3. [목표 우선순위(Priority)]: 작업 시 가장 먼저 고려해야 할 핵심 가치와 순위를 명시할 것.
        4. [실행 임무(Mission)]: 구체적인 결과물 도출을 위한 Few-shot(예시) 템플릿을 최소 2개 이상 포함할 것.
        5. [핵심 준수 규정(Compliance)]: 품질 보장을 위해 반드시 지켜야 할 원칙과 금기 사항을 5가지 이상 기술할 것.

        [출력 규칙]
        - 격조 높고 신뢰감 있는 비즈니스 전문 용어를 사용하라.
        - 마크다운(Markdown)을 활용해 섹션을 일목요연하게 구분하라.
        - 사용자가 즉시 복사하여 업무에 투입할 수 있는 '완성형 도구'의 형태로 제공하라.
        """

        response = active_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_instruction},
                {"role": "user", "content": f"다음 업무 목적을 분석하여 지고의 프롬프트를 설계하라: {user_purpose}"}
            ],
            temperature=0.8
        )
        
        # [해결 포인트] 결과를 문자열이 아닌 '바이트'로 직접 변환하여 전송 (ASCII 간섭 배제)
        result_content = response.choices[0].message.content
        response_data = json.dumps({"result": result_content}, ensure_ascii=False).encode('utf-8')
        
        # 브라우저에 명확하게 UTF-8임을 알려주는 헤더와 함께 Response 객체 반환
        return Response(response_data, status=200, mimetype='application/json; charset=utf-8')

    except Exception as e:
        error_msg = str(e)
        # 에러 발생 시에도 한글이 깨지지 않도록 UTF-8을 명시적으로 선언합니다.
        err_json = json.dumps({"error": error_msg}, ensure_ascii=False).encode('utf-8')
        return Response(err_json, status=500, mimetype='application/json; charset=utf-8')

# ---------------------------------------------------------
# 5. 인증서 및 6. 백업 로직 (디자인 및 멘트 100% 유지)
# ---------------------------------------------------------
@app.route('/generate_certificate', methods=['POST'])
def generate_certificate():
    user_display = "PROMPT MASTER" 
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 24)
    pdf.cell(0, 60, "CERTIFICATE OF COMPLETION", ln=True, align='C')
    pdf.set_font("Arial", size=16)
    pdf.cell(0, 20, "This is to certify that", ln=True, align='C')
    pdf.set_font("Arial", 'B', 30)
    pdf.cell(0, 30, user_display, ln=True, align='C')
    pdf.set_font("Arial", size=14)
    pdf.cell(0, 20, "has successfully mastered the AI Prompting Master Course.", ln=True, align='C')
    
    file_stream = io.BytesIO()
    pdf.output(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name="Official_Certificate.pdf")

@app.route('/backup_all', methods=['POST'])
def backup_all():
    logs = request.json.get('logs', [])
    doc = Document()
    doc.add_heading('My AI Prompt Bible', 0)
    for log in logs:
        doc.add_heading(f"업무 목적: {log['purpose']}", level=1)
        doc.add_paragraph(f"생성일자: {log['date']}")
        doc.add_paragraph(log['content'])
        doc.add_page_break()
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, download_name="My_Prompt_Bible.docx")

@app.route('/download_doc', methods=['POST'])
def download_doc():
    data = request.json
    content, file_format = data.get('content', ''), data.get('format', 'txt')
    filename = f"Prompt_Export_{int(time.time())}"

    if file_format == 'word':
        doc = Document()
        doc.add_heading('AI 실무 프롬프트 결과물', 0)
        doc.add_paragraph(content)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return send_file(file_stream, as_attachment=True, download_name=f"{filename}.docx")
    else:
        # 데이터 손실 방지를 위한 바이트 변환 전송
        file_stream = io.BytesIO(content.encode('utf-8'))
        ext = 'md' if file_format == 'md' else file_format
        return send_file(file_stream, as_attachment=True, download_name=f"{filename}.{ext}")

@app.route('/architect')
def architect():
    data = load_data()
    return render_template('lab.html', patterns=data.get('prompt_patterns', []))

if __name__ == '__main__':
    # [배포용 수정] 
    # 1. host를 '0.0.0.0'으로 설정해야 외부(인터넷)에서 접속이 가능합니다.
    # 2. 포트는 배포 플랫폼(Render 등)이 지정해주는 환경 변수를 우선적으로 따릅니다.
    # 3. 보안을 위해 debug는 False로 변경합니다.
    
    port = int(os.environ.get("PORT", 5000))
    print(f"--- SERVER STARTING ON PORT {port} ---")
    app.run(host='0.0.0.0', port=port, debug=False)